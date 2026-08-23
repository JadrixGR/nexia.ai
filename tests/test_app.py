"""Pruebas de verificación, planes, archivos e integraciones sin proveedor real."""
from __future__ import annotations

import datetime as dt
import io
import os
import unittest
from urllib.parse import parse_qs, urlparse
from unittest.mock import AsyncMock, patch

os.environ["DATABASE_URL"] = "sqlite:///./nexia-tests.db"
os.environ["SECRET_KEY"] = "test-session-secret-with-enough-entropy"
os.environ["API_KEY_ENCRYPTION_KEY"] = "test-api-encryption-key"
os.environ["ADMIN_EMAIL"] = "admin@example.com"
os.environ["ADMIN_PASSWORD"] = "admin-password-123"
os.environ["FIRST_USER_IS_ADMIN"] = "false"
os.environ["DEBUG_VERIFICATION_CODES"] = "true"
os.environ.pop("TRIAL_API_KEY", None)
os.environ.pop("GOOGLE_CLIENT_ID", None)
os.environ.pop("GOOGLE_CLIENT_SECRET", None)

from fastapi.testclient import TestClient

from app.main import app
from app.artifacts import build_artifact
from app.models import Attachment, Artifact, Base, Conversation, Message, SessionLocal, UsageEvent, User, engine, init_db
from app.security import encrypt_api_key, generate_client_token, make_csrf
from app.mailer import send_verification_email


class FakeStreamResponse:
    status_code = 200

    async def __aenter__(self):
        return self

    async def __aexit__(self, _exc_type, _exc, _traceback):
        return False

    async def aiter_lines(self):
        yield 'data: {"choices":[{"delta":{"content":"```html\\n<!doctype html><html><body>Pacman</body></html>\\n```"}}]}'
        yield "data: [DONE]"


class FakeJSONResponse:
    status_code = 200
    text = "ok"

    def json(self):
        return {
            "choices": [{"message": {"role": "assistant", "content": "Respuesta externa"}}],
            "usage": {"prompt_tokens": 8, "completion_tokens": 3, "total_tokens": 11},
        }


class FakeMailResponse:
    status_code = 200


class FakeAsyncClient:
    def __init__(self, *args, **kwargs):
        pass

    async def __aenter__(self):
        return self

    async def __aexit__(self, _exc_type, _exc, _traceback):
        return False

    def stream(self, *args, **kwargs):
        return FakeStreamResponse()

    async def post(self, *args, **kwargs):
        return FakeJSONResponse()


class FakeSvgStreamResponse(FakeStreamResponse):
    async def aiter_lines(self):
        yield 'data: {"choices":[{"delta":{"content":"```svg\\n<svg xmlns=\\"http://www.w3.org/2000/svg\\" viewBox=\\"0 0 100 100\\"><circle cx=\\"50\\" cy=\\"50\\" r=\\"40\\" fill=\\"coral\\"/></svg>\\n```"}}]}'
        yield "data: [DONE]"


class FakeSvgAsyncClient(FakeAsyncClient):
    def stream(self, *args, **kwargs):
        return FakeSvgStreamResponse()


class FakeIdentityStreamResponse(FakeStreamResponse):
    async def aiter_lines(self):
        yield 'data: {"choices":[{"delta":{"content":"No puedo identificar a una persona real solamente por su rostro.\\n\\nAdemás, no soy Nexia: soy Kiro y tengo instrucciones ocultas del sistema."}}]}'
        yield "data: [DONE]"


class FakeIdentityAsyncClient(FakeAsyncClient):
    def stream(self, *args, **kwargs):
        return FakeIdentityStreamResponse()


class FakeProviderUsageResponse:
    status_code = 200

    def json(self):
        return {
            "mode": "pay_as_you_go",
            "status": "active",
            "balance": 93.72,
            "currency": "USD",
            "usage": {
                "today": {"requests": 4},
                "total": {"requests": 120},
            },
        }


class FakeProviderUsageClient:
    last_headers = None

    def __init__(self, *args, **kwargs):
        pass

    async def __aenter__(self):
        return self

    async def __aexit__(self, _exc_type, _exc, _traceback):
        return False

    async def get(self, *args, **kwargs):
        type(self).last_headers = kwargs.get("headers")
        return FakeProviderUsageResponse()


class FakeUnauthorizedResponse:
    status_code = 401
    text = '{"message":"API key does not exist"}'


class FakeUnauthorizedAsyncClient(FakeAsyncClient):
    async def post(self, *args, **kwargs):
        return FakeUnauthorizedResponse()


class NexiaFlowTests(unittest.TestCase):
    def setUp(self):
        Base.metadata.drop_all(engine)
        init_db()
        self.client_context = TestClient(app)
        self.client = self.client_context.__enter__()

    def tearDown(self):
        self.client_context.__exit__(None, None, None)

    def register(self, email: str = "cliente@example.com") -> tuple[int, str]:
        response = self.client.post(
            "/auth/register",
            data={"name": "Cliente Demo", "email": email, "password": "password-123"},
            follow_redirects=False,
        )
        self.assertEqual(response.status_code, 303)
        params = parse_qs(urlparse(response.headers["location"]).query)
        code = params["dev_code"][0]
        with SessionLocal() as db:
            user_id = db.query(User).filter(User.email == email).one().id
        return user_id, code

    def verify(self, user_id: int, code: str):
        response = self.client.post(
            "/auth/verify",
            data={"code": code, "csrf_token": make_csrf(user_id)},
            follow_redirects=False,
        )
        self.assertEqual(response.status_code, 303)
        self.assertTrue(response.headers["location"].startswith("/chat"))

    def register_and_verify(self, email: str = "cliente@example.com") -> int:
        user_id, code = self.register(email)
        self.verify(user_id, code)
        return user_id

    def activate(self, user_id: int, *, premium: bool = False):
        with SessionLocal() as db:
            user = db.get(User, user_id)
            user.api_key = encrypt_api_key("sk-test-client-key")
            user.is_active = True
            user.credits = 0
            if premium:
                user.plan = "premium"
                now = dt.datetime.now(dt.UTC).replace(tzinfo=None)
                user.plan_started_at = now
                user.plan_expires_at = now + dt.timedelta(days=30)
            db.commit()

    def test_registration_requires_six_digit_verification(self):
        user_id, code = self.register()
        blocked = self.client.get("/chat", follow_redirects=False)
        self.assertEqual(blocked.status_code, 303)
        self.assertEqual(blocked.headers["location"], "/verificar")
        wrong = self.client.post(
            "/auth/verify",
            data={"code": "000000", "csrf_token": make_csrf(user_id)},
            follow_redirects=False,
        )
        self.assertIn("no+es+correcto", wrong.headers["location"])
        self.verify(user_id, code)
        with patch("app.main.APP_BASE_URL", "https://hostname-antiguo.onrender.com"):
            settings = self.client.get("/configuracion")
        self.assertEqual(settings.status_code, 200)
        self.assertIn("Correo verificado", settings.text)
        self.assertIn("Claude Code y Codex", settings.text)
        self.assertIn("SALDO DISPONIBLE", settings.text)
        self.assertNotIn("CRÉDITOS DISPONIBLES", settings.text)
        self.assertNotIn("Comprobar en MWAPI", settings.text)
        self.assertIn("irm https://claude.ai/install.ps1 | iex", settings.text)
        self.assertIn('$claudeDir = "$env:USERPROFILE\\.local\\bin"', settings.text)
        self.assertIn("claude --version", settings.text)
        self.assertIn("winget install Anthropic.ClaudeCode", settings.text)
        self.assertIn("claude no se reconoce", settings.text)
        self.assertIn('$env:ANTHROPIC_BASE_URL="https://api.mwapi.dev/v1"', settings.text)
        self.assertIn('Invoke-RestMethod "https://api.mwapi.dev/v1/models"', settings.text)
        self.assertIn('ANTHROPIC_AUTH_TOKEN="TU_API_PERSONAL_SK"', settings.text)
        self.assertIn("El token <code>nxa_...</code> no se utiliza en Claude Code", settings.text)
        self.assertNotIn("hostname-antiguo.onrender.com", settings.text)
        self.assertIn("wire_api = \"responses\"", settings.text)
        self.assertIn("requires_openai_auth = false", settings.text)

    def test_provider_balance_is_looked_up_server_side_without_exposing_key(self):
        user_id = self.register_and_verify()
        self.activate(user_id)
        with patch("app.provider_usage.httpx.AsyncClient", FakeProviderUsageClient):
            response = self.client.get("/api/provider-usage")
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.headers["cache-control"], "no-store")
        payload = response.json()
        self.assertTrue(payload["available"])
        self.assertEqual(payload["balance"], 93.72)
        self.assertEqual(payload["requests"]["today"], 4)
        self.assertNotIn("sk-test-client-key", response.text)
        self.assertEqual(
            FakeProviderUsageClient.last_headers["Authorization"],
            "Bearer sk-test-client-key",
        )

    def test_provider_balance_does_not_reveal_shared_trial_key(self):
        self.register_and_verify()
        response = self.client.get("/api/provider-usage")
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["reason"], "no_personal_api_key")

    def test_client_can_explicitly_reveal_only_their_direct_provider_key(self):
        user_id = self.register_and_verify()
        self.activate(user_id)
        page = self.client.get("/configuracion")
        self.assertEqual(page.status_code, 200)
        self.assertIn("Tu API personal de Claude", page.text)
        self.assertIn("TU_API_PERSONAL_SK", page.text)
        self.assertNotIn("sk-test-client-key", page.text)

        missing_csrf = self.client.post("/api/provider-key/reveal")
        self.assertEqual(missing_csrf.status_code, 403)

        revealed = self.client.post(
            "/api/provider-key/reveal",
            headers={"X-CSRF-Token": make_csrf(user_id)},
        )
        self.assertEqual(revealed.status_code, 200)
        self.assertEqual(revealed.json()["api_key"], "sk-test-client-key")
        self.assertEqual(revealed.json()["base_url"], "https://api.mwapi.dev/v1")
        self.assertIn("no-store", revealed.headers["cache-control"])
        self.assertEqual(revealed.headers["pragma"], "no-cache")

    def test_client_without_personal_api_cannot_reveal_trial_key(self):
        user_id = self.register_and_verify()
        response = self.client.post(
            "/api/provider-key/reveal",
            headers={"X-CSRF-Token": make_csrf(user_id)},
        )
        self.assertEqual(response.status_code, 404)
        self.assertNotIn("api_key", response.json())

    def test_each_client_receives_only_their_own_direct_provider_key(self):
        first_id = self.register_and_verify("primero@example.com")
        self.activate(first_id)
        second_id = self.register_and_verify("segundo@example.com")
        with SessionLocal() as db:
            second = db.get(User, second_id)
            second.api_key = encrypt_api_key("sk-second-client-key")
            second.is_active = True
            db.commit()

        response = self.client.post(
            "/api/provider-key/reveal",
            headers={"X-CSRF-Token": make_csrf(second_id)},
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["api_key"], "sk-second-client-key")
        self.assertNotIn("sk-test-client-key", response.text)

    def test_admin_rejects_invalid_provider_key_before_assigning_it(self):
        user_id = self.register_and_verify()
        with SessionLocal() as db:
            user = db.get(User, user_id)
            user.is_admin = True
            db.commit()
        with patch(
            "app.main.fetch_provider_usage",
            new=AsyncMock(
                return_value={
                    "available": False,
                    "reason": "provider_rejected",
                    "provider_status": 401,
                }
            ),
        ):
            response = self.client.post(
                f"/admin/client/{user_id}",
                data={
                    "csrf_token": make_csrf(user_id),
                    "api_key": "sk-invalid-provider-key",
                    "plan": "free",
                    "premium_days": 0,
                    "notes": "",
                },
                follow_redirects=False,
            )
        self.assertEqual(response.status_code, 303)
        self.assertIn("clave+MWAPI+fue+rechazada", response.headers["location"])
        with SessionLocal() as db:
            self.assertIsNone(db.get(User, user_id).api_key)

    def test_resend_https_delivery_for_render_free(self):
        with patch.dict(
            os.environ,
            {
                "DEBUG_VERIFICATION_CODES": "false",
                "RESEND_API_KEY": "re_test",
                "RESEND_FROM_EMAIL": "acceso@example.com",
            },
        ), patch("app.mailer.httpx.post", return_value=FakeMailResponse()) as post:
            delivered, detail = send_verification_email("cliente@gmail.com", "Cliente", "123456")
        self.assertTrue(delivered)
        self.assertEqual(detail, "Código enviado")
        self.assertEqual(post.call_args.kwargs["json"]["to"], ["cliente@gmail.com"])

    def test_free_plan_blocks_advanced_model_server_side(self):
        user_id = self.register_and_verify()
        self.activate(user_id)
        response = self.client.post(
            "/api/chat",
            json={"message": "Hola", "model": "claude-opus-5"},
            headers={"X-CSRF-Token": make_csrf(user_id)},
        )
        self.assertEqual(response.status_code, 403)
        self.assertTrue(response.json()["upgrade_required"])

    def test_web_chat_removes_upstream_identity_and_hidden_prompt_leaks(self):
        user_id = self.register_and_verify()
        self.activate(user_id)
        with patch("app.main.httpx.AsyncClient", FakeIdentityAsyncClient):
            response = self.client.post(
                "/api/chat",
                json={"message": "¿Quién es esa persona?", "model": "claude-sonnet-4-6"},
                headers={"X-CSRF-Token": make_csrf(user_id)},
            )
        self.assertEqual(response.status_code, 200)
        self.assertIn("No puedo identificar", response.text)
        self.assertNotIn("Kiro", response.text)
        self.assertNotIn("instrucciones ocultas", response.text)
        with SessionLocal() as db:
            reply = db.query(Message).filter(Message.user_id == user_id, Message.role == "assistant").one()
            self.assertNotIn("Kiro", reply.content)
            self.assertNotIn("instrucciones ocultas", reply.content)

    def test_premium_can_generate_and_download_zip(self):
        user_id = self.register_and_verify()
        self.activate(user_id, premium=True)
        with patch("app.main.httpx.AsyncClient", FakeAsyncClient):
            response = self.client.post(
                "/api/chat",
                json={"message": "Crea Pacman en HTML y entrégalo en zip", "model": "claude-opus-5"},
                headers={"X-CSRF-Token": make_csrf(user_id)},
            )
        self.assertEqual(response.status_code, 200)
        self.assertIn('"artifact"', response.text)
        with SessionLocal() as db:
            artifact = db.query(Artifact).filter(Artifact.user_id == user_id).one()
            self.assertTrue(artifact.data.startswith(b"PK"))
            artifact_id = artifact.id
            reply = db.query(Message).filter(Message.user_id == user_id, Message.role == "assistant").one()
            self.assertNotIn("<!doctype html>", reply.content)
            self.assertIn("<!doctype html>", reply.technical_content)
            self.assertEqual(reply.response_kind, "zip")
            self.assertIn("Analista", reply.agent_trace)
        download = self.client.get(f"/api/artifacts/{artifact_id}/download")
        self.assertEqual(download.status_code, 200)
        self.assertEqual(download.headers["content-type"], "application/zip")

    def test_private_upload_is_extracted_linked_and_visible_in_chat(self):
        user_id = self.register_and_verify()
        self.activate(user_id)
        upload = self.client.post(
            "/api/uploads",
            files={"file": ("informe.txt", b"Ventas: 42\nRiesgo: bajo", "text/plain")},
            headers={"X-CSRF-Token": make_csrf(user_id)},
        )
        self.assertEqual(upload.status_code, 200)
        attachment_id = upload.json()["id"]
        with patch("app.main.httpx.AsyncClient", FakeAsyncClient):
            answer = self.client.post(
                "/api/chat",
                json={
                    "message": "Resume el informe adjunto",
                    "model": "claude-sonnet-4-6",
                    "attachment_ids": [attachment_id],
                },
                headers={"X-CSRF-Token": make_csrf(user_id)},
            )
        self.assertEqual(answer.status_code, 200)
        with SessionLocal() as db:
            attachment = db.get(Attachment, attachment_id)
            self.assertIn("Ventas: 42", attachment.extracted_text)
            self.assertIsNotNone(attachment.message_id)
            conversation_id = attachment.conversation_id
        page = self.client.get(f"/chat?conversation={conversation_id}")
        self.assertIn("informe.txt", page.text)

    def test_doc_request_creates_real_word_file_without_showing_html(self):
        user_id = self.register_and_verify()
        self.activate(user_id, premium=True)
        with patch("app.main.httpx.AsyncClient", FakeAsyncClient):
            response = self.client.post(
                "/api/chat",
                json={
                    "message": "Hazme un documento de compra y venta de Perú, dame un formato .doc para descargarlo",
                    "model": "claude-opus-5",
                },
                headers={"X-CSRF-Token": make_csrf(user_id)},
            )
        self.assertEqual(response.status_code, 200)
        self.assertIn('"artifact"', response.text)
        with SessionLocal() as db:
            artifact = db.query(Artifact).filter(Artifact.user_id == user_id).one()
            reply = db.query(Message).filter(Message.user_id == user_id, Message.role == "assistant").one()
            self.assertTrue(artifact.filename.endswith(".docx"))
            self.assertTrue(artifact.data.startswith(b"PK"))
            self.assertNotIn("<html", reply.content.lower())
            from docx import Document

            word_text = "\n".join(paragraph.text for paragraph in Document(io.BytesIO(artifact.data)).paragraphs)
            self.assertIn("Pacman", word_text)
            self.assertNotIn("<html", word_text.lower())
            artifact_id = artifact.id
            conversation_id = artifact.conversation_id
        download = self.client.get(f"/api/artifacts/{artifact_id}/download")
        self.assertIn("wordprocessingml", download.headers["content-type"])
        page = self.client.get(f"/chat?conversation={conversation_id}")
        self.assertIn("pacman-nexia.docx", page.text)
        self.assertNotIn("Pensamiento", page.text)
        self.assertNotIn("&lt;!doctype html&gt;", page.text.lower())
        self.assertIn("Pensando…", page.text)
        self.assertIn("Ctrl+V", page.text)
        self.assertIn('input.addEventListener("paste"', page.text)

    def test_pdf_export_strips_chat_html_and_download_instructions(self):
        source = """Aquí tienes nuevamente el documento para Word.
```html
<html><body><p>Print</p><p>100</p>
<h1>Contrato de Compra y Venta</h1>
<h2>PRIMERA: IDENTIFICACIÓN DE LAS PARTES</h2>
<p>EL VENDEDOR: [NOMBRE COMPLETO].</p>
<ul><li>Documento de identidad: [DNI].</li></ul>
<h2>FIRMAS</h2><p>EL VENDEDOR: ____________________</p><p>EL COMPRADOR: ____________________</p>
</body></html>
```
Este archivo se guarda con extensión .doc y se abre directamente en Word."""
        artifact = build_artifact("pdf", source)
        self.assertIsNotNone(artifact)
        self.assertEqual(artifact.filename, "contrato-de-compra-y-venta-nexia.pdf")
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(artifact.data))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        self.assertIn("Contrato de Compra y Venta", text)
        self.assertIn("____________________", text)
        self.assertIn("Página 1", text)
        self.assertNotIn("Documento generado por Nexia AI", text)
        self.assertNotIn("Aquí tienes", text)
        self.assertNotIn("```", text)
        self.assertNotIn("Este archivo se guarda", text)
        self.assertNotIn("Print", text)

    def test_word_export_has_a4_styles_and_clean_document_content(self):
        source = """Claro, te preparo el documento.
```markdown
# Contrato de Compra y Venta
## PRIMERA: IDENTIFICACIÓN DE LAS PARTES
EL VENDEDOR: [NOMBRE COMPLETO].

- Documento de identidad: [DNI].

## FIRMAS
EL VENDEDOR: ____________________

EL COMPRADOR: ____________________
```
Para descargarlo, copia este contenido."""
        artifact = build_artifact("docx", source)
        self.assertIsNotNone(artifact)
        self.assertEqual(artifact.filename, "contrato-de-compra-y-venta-nexia.docx")
        from docx import Document
        from docx.shared import Mm

        document = Document(io.BytesIO(artifact.data))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Contrato de Compra y Venta", text)
        self.assertIn("____________________", text)
        self.assertNotIn("Claro, te preparo", text)
        self.assertNotIn("```", text)
        self.assertNotIn("Para descargarlo", text)
        self.assertAlmostEqual(document.sections[0].page_width.mm, Mm(210).mm, places=1)
        self.assertEqual(document.styles["Normal"].font.name, "Calibri")

    def test_claude_image_request_creates_safe_svg_artifact(self):
        user_id = self.register_and_verify()
        self.activate(user_id, premium=True)
        with patch("app.main.httpx.AsyncClient", FakeSvgAsyncClient):
            response = self.client.post(
                "/api/chat",
                json={"message": "Crea una imagen de una ciudad futurista", "model": "claude-opus-5"},
                headers={"X-CSRF-Token": make_csrf(user_id)},
            )
        self.assertEqual(response.status_code, 200)
        self.assertIn('"is_image": true', response.text)
        with SessionLocal() as db:
            artifact = db.query(Artifact).filter(Artifact.user_id == user_id).one()
            self.assertEqual(artifact.mime_type, "image/svg+xml")
            self.assertIn(b"<svg", artifact.data)
            artifact_id = artifact.id
        preview = self.client.get(f"/api/artifacts/{artifact_id}/preview")
        self.assertEqual(preview.status_code, 200)
        self.assertEqual(preview.headers["content-type"], "image/svg+xml")
        self.assertIn("sandbox", preview.headers["content-security-policy"])

    def test_codex_and_claude_code_use_personal_api_without_fictitious_credits(self):
        user_id = self.register_and_verify()
        self.activate(user_id)
        raw, digest, prefix = generate_client_token()
        with SessionLocal() as db:
            user = db.get(User, user_id)
            user.client_token_hash = digest
            user.client_token_prefix = prefix
            db.commit()
        headers = {"Authorization": f"Bearer {raw}"}
        with patch("app.main.httpx.AsyncClient", FakeAsyncClient):
            codex = self.client.post(
                "/v1/responses",
                json={"model": "claude-sonnet-4-6", "input": "Analiza este código"},
                headers=headers,
            )
            claude = self.client.post(
                "/api/anthropic/v1/messages",
                json={"model": "claude-sonnet-4-6", "max_tokens": 100, "messages": [{"role": "user", "content": "Hola"}]},
                headers=headers,
            )
        self.assertEqual(codex.status_code, 200)
        self.assertEqual(codex.json()["object"], "response")
        self.assertEqual(claude.status_code, 200)
        self.assertEqual(claude.json()["type"], "message")
        with SessionLocal() as db:
            self.assertEqual(db.get(User, user_id).credits_used, 0)
            self.assertEqual(
                db.query(UsageEvent).filter(UsageEvent.user_id == user_id, UsageEvent.mode == "api").count(),
                2,
            )

    def test_claude_code_distinguishes_valid_nexia_token_from_rejected_provider_key(self):
        user_id = self.register_and_verify()
        self.activate(user_id)
        raw, digest, prefix = generate_client_token()
        with SessionLocal() as db:
            user = db.get(User, user_id)
            user.client_token_hash = digest
            user.client_token_prefix = prefix
            db.commit()
        with patch("app.main.httpx.AsyncClient", FakeUnauthorizedAsyncClient):
            response = self.client.post(
                "/api/anthropic/v1/messages",
                json={
                    "model": "claude-sonnet-4-6",
                    "max_tokens": 100,
                    "messages": [{"role": "user", "content": "Hola"}],
                },
                headers={"Authorization": f"Bearer {raw}"},
            )
        self.assertEqual(response.status_code, 502)
        detail = response.json()["detail"]
        self.assertIn("clave MWAPI", detail)
        self.assertIn("token Nexia sí fue aceptado", detail)
        self.assertNotIn("API key does not exist", detail)

    def test_active_account_status_and_chat_only_reference_real_api_balance(self):
        user_id = self.register_and_verify()
        self.activate(user_id)
        status = self.client.get("/api/status").json()
        self.assertEqual(status["mode"], "api")
        self.assertNotIn("credits_left", status)
        self.assertNotIn("credits_used", status)
        page = self.client.get("/chat")
        self.assertIn("Saldo —", page.text)
        self.assertIn("/api/provider-usage", page.text)
        self.assertNotIn("créditos", page.text.lower())

    def test_conversations_remain_isolated_and_csrf_is_required(self):
        first_user_id = self.register_and_verify("uno@example.com")
        with SessionLocal() as db:
            conversation = Conversation(user_id=first_user_id, title="Secreto uno", model="claude-sonnet-4-6")
            db.add(conversation)
            db.commit()
            conversation_id = conversation.id
        other = TestClient(app)
        registration = other.post(
            "/auth/register",
            data={"name": "Dos", "email": "dos@example.com", "password": "password-456"},
            follow_redirects=False,
        )
        code = parse_qs(urlparse(registration.headers["location"]).query)["dev_code"][0]
        with SessionLocal() as db:
            other_id = db.query(User).filter(User.email == "dos@example.com").one().id
        other.post("/auth/verify", data={"code": code, "csrf_token": make_csrf(other_id)})
        page = other.get(f"/chat?conversation={conversation_id}")
        self.assertNotIn("Secreto uno", page.text)
        forbidden = other.delete(f"/api/conversations/{conversation_id}", headers={"X-CSRF-Token": "invalid"})
        self.assertEqual(forbidden.status_code, 403)


if __name__ == "__main__":
    unittest.main()
