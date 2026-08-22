"""Crea archivos reales a partir de una respuesta y los guarda como binarios."""
from __future__ import annotations

import io
import re
import unicodedata
import zipfile
from dataclasses import dataclass
from html import escape, unescape


@dataclass
class BuiltArtifact:
    filename: str
    mime_type: str
    data: bytes


def requested_kind(prompt: str) -> str | None:
    text = prompt.lower()
    if re.search(r"\b(zip|comprimid[oa]|empaquet\w*)\b", text):
        return "zip"
    if re.search(r"\b(pdf)\b", text):
        return "pdf"
    if re.search(r"\b(excel|xlsx|hoja de c[aá]lculo)\b", text):
        return "xlsx"
    if re.search(r"\bword\b", text) or re.search(r"(?<![\w])\.?docx?(?![\w])", text):
        return "docx"
    if re.search(r"\b(html)\b", text) and re.search(r"\b(archivo|formato|entrega|descarga|crea|haz)\w*\b", text):
        return "html"
    return None


def _code_blocks(answer: str) -> list[tuple[str, str]]:
    return [(lang.lower(), code.strip()) for lang, code in re.findall(r"```([\w.+-]*)\s*\n(.*?)```", answer, re.S)]


def _best_html(answer: str) -> str:
    for language, code in _code_blocks(answer):
        if language in {"html", "htm"} or "<!doctype html" in code.lower() or "<html" in code.lower():
            return code
    return answer if "<html" in answer.lower() else f"<!doctype html><html><meta charset='utf-8'><body><pre>{escape(answer)}</pre></body></html>"


def _document_source(answer: str) -> str:
    """Selecciona el documento y descarta la conversación que lo envuelve."""
    candidates = [
        code for language, code in _code_blocks(answer)
        if language in {"", "html", "htm", "markdown", "md", "text", "txt"}
    ]
    if candidates:
        return max(candidates, key=len)
    return answer


def _document_text(answer: str) -> str:
    """Convierte HTML accidental del modelo a texto estructurado para Word/PDF."""
    clean = _document_source(answer).replace("\r\n", "\n").replace("\r", "\n")
    clean = re.sub(r"<(script|style)\b[^>]*>.*?</\1>", "", clean, flags=re.I | re.S)
    clean = re.sub(r"<!--.*?-->", "", clean, flags=re.S)
    clean = re.sub(r"<h1\b[^>]*>", "\n# ", clean, flags=re.I)
    clean = re.sub(r"<h2\b[^>]*>", "\n## ", clean, flags=re.I)
    clean = re.sub(r"<h[3-6]\b[^>]*>", "\n### ", clean, flags=re.I)
    clean = re.sub(r"<li\b[^>]*>", "\n- ", clean, flags=re.I)
    clean = re.sub(r"<(br|/p|/div|/li|/h[1-6]|/tr)\b[^>]*>\s*", "\n", clean, flags=re.I)
    clean = re.sub(r"<(p|div|tr)\b[^>]*>", "\n", clean, flags=re.I)
    clean = re.sub(r"</t[dh]>\s*", " | ", clean, flags=re.I)
    clean = re.sub(r"<[^>]+>", "", clean)
    clean = unescape(clean).replace("\xa0", " ")
    clean = re.sub(r"^\s*```[\w.+-]*\s*$", "", clean, flags=re.M)
    clean = clean.replace("■", "[ ]").replace("☐", "[ ]")
    clean = re.sub(r"[ \t]+$", "", clean, flags=re.M)
    lines = clean.splitlines()
    while lines and (not lines[0].strip() or re.fullmatch(r"(?i)(print|100)", lines[0].strip())):
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    clean = "\n".join(lines)
    clean = re.sub(
        r"(?ims)\n+(?:este archivo se guarda|para (?:guardarlo|descargarlo)|copia (?:este|el) contenido|"
        r"si deseas,? (?:también )?puedo).*$",
        "",
        clean,
    )
    clean = re.sub(r"\n{3,}", "\n\n", clean).strip()

    # Elimina títulos consecutivos repetidos (por ejemplo texto plano + '# Título').
    compact = [line for line in clean.splitlines() if line.strip()]
    if len(compact) >= 2:
        first = re.sub(r"^[#*\s]+|[#*\s]+$", "", compact[0]).strip().casefold()
        second = re.sub(r"^[#*\s]+|[#*\s]+$", "", compact[1]).strip().casefold()
        if first == second:
            clean = re.sub(r"^\s*" + re.escape(compact[0]) + r"\s*\n+", "", clean, count=1)
    return clean


@dataclass
class DocumentBlock:
    kind: str
    text: str
    level: int = 0


def _plain_inline(value: str) -> str:
    value = re.sub(r"!\[([^]]*)]\([^)]*\)", r"\1", value)
    value = re.sub(r"\[([^]]+)]\([^)]*\)", r"\1", value)
    value = re.sub(r"(?<!\\)([*~`]{1,3})(.*?)\1", r"\2", value)
    value = re.sub(r"(?<!_)__([^_\n]+)__(?!_)", r"\1", value)
    value = re.sub(r"(?<!_)_([^_\n]+)_(?!_)", r"\1", value)
    return re.sub(r"\\([#*_[\]()`])", r"\1", value).strip()


def _document_blocks(answer: str) -> list[DocumentBlock]:
    text = _document_text(answer)
    blocks: list[DocumentBlock] = []
    paragraph: list[str] = []

    def flush() -> None:
        if paragraph:
            value = _plain_inline(" ".join(part.strip() for part in paragraph if part.strip()))
            if value:
                blocks.append(DocumentBlock("paragraph", value))
            paragraph.clear()

    for raw_line in [*text.splitlines(), ""]:
        line = raw_line.strip()
        if not line:
            flush()
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            flush()
            blocks.append(DocumentBlock("heading", _plain_inline(heading.group(2)), min(len(heading.group(1)), 3)))
            continue
        bullet = re.match(r"^(?:[-*+]\s+|\[([ xX])\]\s*)(.+)$", line)
        if bullet:
            flush()
            prefix = "[x] " if bullet.group(1) and bullet.group(1).lower() == "x" else ""
            blocks.append(DocumentBlock("bullet", prefix + _plain_inline(bullet.group(2))))
            continue
        numbered = re.match(r"^(\d+[.)])\s+(.+)$", line)
        if numbered:
            flush()
            blocks.append(DocumentBlock("numbered", f"{numbered.group(1)} {_plain_inline(numbered.group(2))}"))
            continue
        if re.match(r"^(?:CL[ÁA]USULA\s+)?(?:PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|S[ÉE]PTIMA|OCTAVA|NOVENA|D[ÉE]CIMA)\b", line, re.I):
            flush()
            blocks.append(DocumentBlock("heading", _plain_inline(line), 2))
            continue
        paragraph.append(line)

    if not blocks:
        blocks.append(DocumentBlock("paragraph", "Documento"))
    return blocks


def _document_title(blocks: list[DocumentBlock]) -> tuple[str, int | None]:
    for index, block in enumerate(blocks):
        if block.kind == "heading" and block.level == 1:
            return block.text, index
    first = blocks[0]
    if first.kind in {"paragraph", "heading"} and len(first.text) <= 120:
        return first.text, 0
    return "Documento", None


def _ascii_stem(title: str) -> str:
    value = unicodedata.normalize("NFKD", title).encode("ascii", "ignore").decode("ascii").lower()
    value = re.sub(r"[^a-z0-9]+", "-", value).strip("-")
    return (value[:64].rstrip("-") or "documento") + "-nexia"


def _best_svg(answer: str) -> str:
    candidates = [code for language, code in _code_blocks(answer) if language == "svg"]
    source = candidates[0] if candidates else answer
    match = re.search(r"<svg\b[^>]*>.*?</svg>", source, re.I | re.S)
    if not match:
        raise ValueError("El modelo no devolvió un SVG válido.")
    svg = match.group(0)
    svg = re.sub(r"<(script|foreignObject)\b[^>]*>.*?</\1>", "", svg, flags=re.I | re.S)
    svg = re.sub(r"\s+on[a-z]+\s*=\s*(['\"]).*?\1", "", svg, flags=re.I | re.S)
    svg = re.sub(r"\s+(?:xlink:)?href\s*=\s*(['\"])(?!#).*?\1", "", svg, flags=re.I | re.S)
    return svg


def _zip(answer: str) -> BuiltArtifact:
    extensions = {
        "html": "html", "htm": "html", "css": "css", "javascript": "js", "js": "js",
        "typescript": "ts", "ts": "ts", "python": "py", "py": "py", "json": "json",
        "markdown": "md", "md": "md", "sql": "sql", "svg": "svg", "xml": "xml",
    }
    names = {"html": "index.html", "css": "styles.css", "js": "script.js", "py": "app.py", "md": "README.md"}
    used: set[str] = set()
    memory = io.BytesIO()
    with zipfile.ZipFile(memory, "w", zipfile.ZIP_DEFLATED) as archive:
        blocks = _code_blocks(answer)
        for number, (language, code) in enumerate(blocks, 1):
            extension = extensions.get(language, "txt")
            filename = names.get(extension, f"archivo-{number}.{extension}")
            if filename in used:
                stem, dot, suffix = filename.rpartition(".")
                filename = f"{stem}-{number}.{suffix}" if dot else f"{filename}-{number}"
            used.add(filename)
            archive.writestr(filename, code.encode("utf-8"))
        if not blocks:
            archive.writestr("respuesta.md", answer.encode("utf-8"))
        elif "README.md" not in used:
            archive.writestr("README.md", "# Archivo generado por Nexia AI\n\nAbre `index.html` si el proyecto es web.\n")
    return BuiltArtifact("proyecto-nexia.zip", "application/zip", memory.getvalue())


def _pdf(answer: str) -> BuiltArtifact:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import KeepTogether, Paragraph, SimpleDocTemplate, Spacer

    blocks = _document_blocks(answer)
    title, title_index = _document_title(blocks)
    memory = io.BytesIO()
    doc = SimpleDocTemplate(
        memory,
        pagesize=A4,
        rightMargin=22 * mm,
        leftMargin=22 * mm,
        topMargin=20 * mm,
        bottomMargin=22 * mm,
        title=title,
        author="Nexia AI",
    )
    styles = getSampleStyleSheet()
    ink = HexColor("#202124")
    muted = HexColor("#687078")
    accent = HexColor("#2E5D78")
    title_style = ParagraphStyle(
        "NexiaTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=18,
        leading=22, textColor=ink, alignment=TA_CENTER, spaceAfter=7 * mm,
    )
    heading_one = ParagraphStyle(
        "NexiaHeading1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=13,
        leading=16, textColor=accent, spaceBefore=6 * mm, spaceAfter=2.5 * mm, keepWithNext=True,
    )
    heading_two = ParagraphStyle(
        "NexiaHeading2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11.5,
        leading=14, textColor=ink, spaceBefore=4.5 * mm, spaceAfter=2 * mm, keepWithNext=True,
    )
    heading_three = ParagraphStyle(
        "NexiaHeading3", parent=heading_two, fontSize=10.5, leading=13, textColor=muted,
        spaceBefore=3.5 * mm, spaceAfter=1.5 * mm,
    )
    body = ParagraphStyle(
        "NexiaBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5,
        leading=15, textColor=ink, alignment=TA_JUSTIFY, spaceAfter=3.2 * mm,
        allowWidows=0, allowOrphans=0,
    )
    bullet = ParagraphStyle(
        "NexiaBullet", parent=body, alignment=TA_LEFT, leftIndent=7 * mm,
        firstLineIndent=-4 * mm, spaceAfter=1.8 * mm,
    )

    def flowables(block: DocumentBlock) -> list:
        value = escape(block.text)
        if block.kind == "heading":
            style = {1: heading_one, 2: heading_two}.get(block.level, heading_three)
            return [Paragraph(value, style)]
        if block.kind == "bullet":
            marker = "[x]" if block.text.startswith("[x] ") else "&bull;"
            text = block.text[4:] if block.text.startswith("[x] ") else block.text
            return [Paragraph(f"{marker}&nbsp;&nbsp;{escape(text)}", bullet)]
        if block.kind == "numbered":
            return [Paragraph(value, bullet)]
        return [Paragraph(value, body)]

    story = [Paragraph(escape(title), title_style)]
    signature_group: list = []
    in_signatures = False
    for index, block in enumerate(blocks):
        if index == title_index:
            continue
        if block.kind == "heading" and re.search(r"\bfirmas?\b", block.text, re.I):
            in_signatures = True
        target = signature_group if in_signatures else story
        target.extend(flowables(block))
    if signature_group:
        story.extend([Spacer(1, 3 * mm), KeepTogether(signature_group)])

    def page_footer(canvas, built_doc) -> None:
        canvas.saveState()
        width, _height = A4
        y = 11 * mm
        canvas.setStrokeColor(HexColor("#D8DDE2"))
        canvas.setLineWidth(0.35)
        canvas.line(22 * mm, y + 5 * mm, width - 22 * mm, y + 5 * mm)
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(muted)
        canvas.drawString(22 * mm, y, "Nexia AI")
        canvas.drawRightString(width - 22 * mm, y, f"Página {built_doc.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=page_footer, onLaterPages=page_footer)
    return BuiltArtifact(f"{_ascii_stem(title)}.pdf", "application/pdf", memory.getvalue())


def _docx(answer: str) -> BuiltArtifact:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt, RGBColor

    blocks = _document_blocks(answer)
    title, title_index = _document_title(blocks)
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.right_margin = Mm(22)
    section.bottom_margin = Mm(22)
    section.left_margin = Mm(22)
    section.header_distance = Mm(12.5)
    section.footer_distance = Mm(12.5)

    def set_style_font(style, name: str, size: float, color: str = "202124", bold: bool | None = None) -> None:
        style.font.name = name
        style._element.rPr.rFonts.set(qn("w:ascii"), name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        if bold is not None:
            style.font.bold = bold

    normal = document.styles["Normal"]
    set_style_font(normal, "Calibri", 11, "202124")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, "2E74B5", 14, 8),
        "Heading 2": (13, "2E74B5", 11, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        set_style_font(style, "Calibri", size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Nexia Title" not in [style.name for style in document.styles]:
        title_style = document.styles.add_style("Nexia Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        title_style = document.styles["Nexia Title"]
    set_style_font(title_style, "Calibri", 20, "202124", True)
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(20)
    title_style.paragraph_format.keep_with_next = True

    list_style = document.styles["List Bullet"]
    set_style_font(list_style, "Calibri", 11, "202124")
    list_style.paragraph_format.left_indent = Mm(9.5)
    list_style.paragraph_format.first_line_indent = Mm(-4.8)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.25

    title_paragraph = document.add_paragraph(style=title_style)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.add_run(title)

    signature_started = False
    signature_paragraphs = []
    for index, block in enumerate(blocks):
        if index == title_index:
            continue
        if block.kind == "heading":
            paragraph = document.add_heading(block.text, level=max(1, min(block.level, 3)))
            if re.search(r"\bfirmas?\b", block.text, re.I):
                signature_started = True
        elif block.kind == "bullet":
            text = block.text[4:] if block.text.startswith("[x] ") else block.text
            prefix = "[x] " if block.text.startswith("[x] ") else ""
            paragraph = document.add_paragraph(prefix + text, style="List Bullet")
        elif block.kind == "numbered":
            paragraph = document.add_paragraph(block.text)
            paragraph.paragraph_format.left_indent = Mm(9.5)
            paragraph.paragraph_format.first_line_indent = Mm(-4.8)
        else:
            paragraph = document.add_paragraph(block.text)
        if block.kind == "paragraph":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if signature_started:
            signature_paragraphs.append(paragraph)

    for paragraph in signature_paragraphs[:-1]:
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_paragraph.add_run("Nexia AI  ·  Página ")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string("687078")
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    footer_run._r.append(field_begin)
    footer_run._r.append(instruction)
    footer_run._r.append(field_end)

    document.core_properties.title = title
    document.core_properties.author = "Nexia AI"
    document.core_properties.subject = "Documento generado y editable"
    memory = io.BytesIO()
    document.save(memory)
    return BuiltArtifact(
        f"{_ascii_stem(title)}.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        memory.getvalue(),
    )


def _xlsx(answer: str) -> BuiltArtifact:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Nexia"
    table_lines = [line for line in answer.splitlines() if line.strip().startswith("|") and line.strip().endswith("|")]
    rows: list[list[str]] = []
    for line in table_lines:
        values = [value.strip() for value in line.strip().strip("|").split("|")]
        if values and not all(re.fullmatch(r":?-{3,}:?", value or "") for value in values):
            rows.append(values)
    if not rows:
        rows = [["Contenido"], *[[line] for line in answer.splitlines() if line.strip()]]
    for row in rows:
        sheet.append(row)
    if rows:
        for cell in sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="C15F3C")
    for column in sheet.columns:
        letter = column[0].column_letter
        sheet.column_dimensions[letter].width = min(max(len(str(cell.value or "")) for cell in column) + 3, 55)
    sheet.freeze_panes = "A2"
    memory = io.BytesIO()
    workbook.save(memory)
    return BuiltArtifact("datos-nexia.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", memory.getvalue())


def build_artifact(kind: str | None, answer: str) -> BuiltArtifact | None:
    if not kind:
        return None
    if kind == "zip":
        return _zip(answer)
    if kind == "pdf":
        return _pdf(answer)
    if kind == "docx":
        return _docx(answer)
    if kind == "xlsx":
        return _xlsx(answer)
    if kind == "html":
        return BuiltArtifact("proyecto-nexia.html", "text/html; charset=utf-8", _best_html(answer).encode("utf-8"))
    if kind == "svg":
        return BuiltArtifact("imagen-nexia.svg", "image/svg+xml", _best_svg(answer).encode("utf-8"))
    return None
