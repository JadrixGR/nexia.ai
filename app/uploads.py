"""Validación y extracción local de archivos subidos por cada cliente."""
from __future__ import annotations

import csv
import io
import json
from pathlib import Path

MAX_UPLOAD_BYTES = 15 * 1024 * 1024
MAX_EXTRACTED_CHARS = 80_000

ALLOWED_EXTENSIONS = {
    ".pdf", ".docx", ".xlsx", ".csv", ".txt", ".md", ".json", ".html",
    ".py", ".js", ".ts", ".css", ".xml", ".png", ".jpg", ".jpeg", ".webp", ".gif",
}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
IMAGE_MIMES = {"image/png", "image/jpeg", "image/webp", "image/gif"}


def safe_filename(value: str) -> str:
    name = Path(value or "archivo").name[:180]
    return "".join(character if character.isalnum() or character in " ._-()" else "-" for character in name) or "archivo"


def validate_upload(filename: str, mime_type: str, data: bytes) -> tuple[str, str]:
    clean_name = safe_filename(filename)
    extension = Path(clean_name).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise ValueError("Formato no permitido. Usa PDF, DOCX, XLSX, CSV, texto, código o una imagen.")
    if not data:
        raise ValueError("El archivo está vacío.")
    if len(data) > MAX_UPLOAD_BYTES:
        raise ValueError("El archivo supera el máximo de 15 MB.")
    clean_mime = (mime_type or "application/octet-stream").split(";")[0].strip().lower()
    if extension in IMAGE_EXTENSIONS and clean_mime not in IMAGE_MIMES:
        clean_mime = {".png": "image/png", ".webp": "image/webp", ".gif": "image/gif"}.get(extension, "image/jpeg")
    return clean_name, clean_mime


def is_image(filename: str, mime_type: str) -> bool:
    return Path(filename).suffix.lower() in IMAGE_EXTENSIONS or mime_type in IMAGE_MIMES


def extract_text(filename: str, mime_type: str, data: bytes) -> str:
    extension = Path(filename).suffix.lower()
    text = ""
    try:
        if extension == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            text = "\n\n".join((page.extract_text() or "") for page in reader.pages[:100])
        elif extension == ".docx":
            from docx import Document

            document = Document(io.BytesIO(data))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                paragraphs.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
            text = "\n".join(paragraphs)
        elif extension == ".xlsx":
            from openpyxl import load_workbook

            workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            lines: list[str] = []
            for sheet in workbook.worksheets[:20]:
                lines.append(f"## Hoja: {sheet.title}")
                for row in sheet.iter_rows(max_row=2000, values_only=True):
                    lines.append(" | ".join("" if value is None else str(value) for value in row))
            text = "\n".join(lines)
        elif extension == ".csv":
            decoded = data.decode("utf-8-sig", "replace")
            rows = list(csv.reader(io.StringIO(decoded)))
            text = "\n".join(" | ".join(row) for row in rows[:5000])
        elif extension == ".json":
            decoded = data.decode("utf-8-sig", "replace")
            text = json.dumps(json.loads(decoded), ensure_ascii=False, indent=2)
        elif not is_image(filename, mime_type):
            text = data.decode("utf-8-sig", "replace")
    except Exception as exc:
        raise ValueError(f"No se pudo leer el contenido de {filename}: {exc.__class__.__name__}.") from exc
    return text.strip()[:MAX_EXTRACTED_CHARS]
